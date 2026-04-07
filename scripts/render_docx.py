#!/usr/bin/env python3
"""
DOCX 渲染引擎 — 基于 python-docx，投研报告级排版。

JSON 接口与 render_pdf.py 完全一致，同一份数据可同时出 PDF 和 DOCX。

支持主题: cicc / ms / cms / dachen
支持章节: cover, toc, heading, paragraph, table, chart, quote, kpi, disclaimer, pagebreak

用法:
  python render_docx.py --data input.json --output report.docx --theme cicc
"""

import json
import os
import tempfile
from typing import Dict, List, Any, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

try:
    from scripts.render_charts import render_chart, CHART_TYPES
except ImportError:
    from render_charts import render_chart, CHART_TYPES


# ═══════════════════════════════════════════
# 主题配色（与 PDF/PPTX 统一）
# ═══════════════════════════════════════════

THEMES = {
    'cicc': {
        'primary': RGBColor(0x1a, 0x1a, 0x2e),
        'accent': RGBColor(0xc4, 0x1e, 0x3a),
        'text_dark': RGBColor(0x1a, 0x1a, 0x2e),
        'text_body': RGBColor(0x33, 0x33, 0x33),
        'text_secondary': RGBColor(0x66, 0x66, 0x66),
        'table_header_bg': '1a1a2e',
        'table_alt_row': 'f5f5f5',
        'quote_bg': 'fdf2f2',
        'cover_bg': '1a1a2e',
        'positive': RGBColor(0x38, 0xa1, 0x69),
        'negative': RGBColor(0xe5, 0x3e, 0x3e),
        'serif_body': True,
        'font_body': 'Noto Serif CJK SC',
        'font_heading': 'Noto Sans CJK SC',
        'label': '中金公司 CICC',
    },
    'ms': {
        'primary': RGBColor(0x00, 0x2D, 0x72),
        'accent': RGBColor(0x00, 0x78, 0xC8),
        'text_dark': RGBColor(0x00, 0x2D, 0x72),
        'text_body': RGBColor(0x4D, 0x4D, 0x4D),
        'text_secondary': RGBColor(0x80, 0x80, 0x80),
        'table_header_bg': '002D72',
        'table_alt_row': 'F2F2F2',
        'quote_bg': 'EBF0F7',
        'cover_bg': '002D72',
        'positive': RGBColor(0x38, 0xa1, 0x69),
        'negative': RGBColor(0xe5, 0x3e, 0x3e),
        'serif_body': False,
        'font_body': 'Noto Sans CJK SC',
        'font_heading': 'Noto Sans CJK SC',
        'label': 'Morgan Stanley',
    },
    'cms': {
        'primary': RGBColor(0xC1, 0x00, 0x2A),
        'accent': RGBColor(0xE6, 0x00, 0x33),
        'text_dark': RGBColor(0x33, 0x33, 0x33),
        'text_body': RGBColor(0x33, 0x33, 0x33),
        'text_secondary': RGBColor(0x66, 0x66, 0x66),
        'table_header_bg': 'A00020',
        'table_alt_row': 'F2F2F2',
        'quote_bg': 'FDF2F4',
        'cover_bg': 'C1002A',
        'positive': RGBColor(0x38, 0xa1, 0x69),
        'negative': RGBColor(0xe5, 0x3e, 0x3e),
        'serif_body': True,
        'font_body': 'Noto Serif CJK SC',
        'font_heading': 'Noto Sans CJK SC',
        'label': '招商证券 CMS',
    },
    'dachen': {
        'primary': RGBColor(0xb8, 0x14, 0x1d),
        'accent': RGBColor(0x18, 0x5d, 0xb1),
        'text_dark': RGBColor(0x33, 0x33, 0x33),
        'text_body': RGBColor(0x33, 0x33, 0x33),
        'text_secondary': RGBColor(0x66, 0x66, 0x66),
        'table_header_bg': 'b8141d',
        'table_alt_row': 'f6f3f4',
        'quote_bg': 'f6f3f4',
        'cover_bg': 'b8141d',
        'positive': RGBColor(0x38, 0xa1, 0x69),
        'negative': RGBColor(0xe5, 0x3e, 0x3e),
        'serif_body': True,
        'font_body': 'Noto Serif CJK SC',
        'font_heading': 'Noto Sans CJK SC',
        'label': '达晨财智 Fortune Capital',
    },
}


def _hex_to_rgb(hex_str: str) -> RGBColor:
    hex_str = hex_str.lstrip('#')
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def _set_cell_shading(cell, color_hex: str):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_paragraph_spacing(para, before=0, after=0, line=None):
    """设置段落间距"""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


def _add_horizontal_line(doc, color: RGBColor, width_pt: float = 1.5):
    """添加水平装饰线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    # 使用底部边框模拟装饰线
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{int(width_pt * 8)}" w:space="1" '
        f'    w:color="{color[0]:02x}{color[1]:02x}{color[2]:02x}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def _add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()


# ═══════════════════════════════════════════
# 图表生成
# ═══════════════════════════════════════════

def _generate_charts(data: Dict, work_dir: str, theme_name: str) -> Dict[str, str]:
    """生成所有图表 PNG"""
    chart_paths = {}
    for i, chart_def in enumerate(data.get('charts', [])):
        chart_id = chart_def.get('id', f'chart_{i}')
        chart_type = chart_def.get('type', 'bar')
        chart_data = chart_def.get('data', {})
        if 'title' not in chart_data:
            chart_data['title'] = chart_def.get('title', '')

        out_path = os.path.join(work_dir, f'{chart_id}.png')
        try:
            chart_theme = {'cicc': 'cicc', 'ms': 'goldman', 'cms': 'goldman'}.get(theme_name, 'default')
            render_chart(chart_type, chart_data, out_path, chart_theme)
            chart_paths[chart_id] = out_path
        except Exception as e:
            print(f"  WARNING: Chart '{chart_id}' failed: {e}")
    return chart_paths


# ═══════════════════════════════════════════
# 文档构建
# ═══════════════════════════════════════════

def _build_cover(doc: Document, data: Dict, theme: Dict):
    """构建封面页"""
    # 封面用全色背景段落模拟
    title_zh = data.get('title', '研究报告')
    title_en = data.get('title_en', '')
    author = data.get('author', '')
    date = data.get('date', '')
    version = data.get('version', '')

    # 添加多个空行推到页面中部
    for _ in range(6):
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, 0, 0)

    # 装饰线
    _add_horizontal_line(doc, theme['accent'], 2.0)

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_zh)
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = theme['primary']
    run.font.name = theme['font_heading']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), theme['font_heading'])
    _set_paragraph_spacing(p, 20, 20)

    # 装饰线
    _add_horizontal_line(doc, theme['accent'], 2.0)

    # 英文副标题
    if title_en:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_en)
        run.font.size = Pt(14)
        run.font.color.rgb = theme['text_secondary']
        _set_paragraph_spacing(p, 20, 10)

    # 作者
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(author)
        run.font.size = Pt(12)
        run.font.color.rgb = theme['text_secondary']
        _set_paragraph_spacing(p, 30, 4)

    # 日期 + 版本
    date_str = date
    if version:
        date_str += f' · {version}'
    if date_str:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(date_str)
        run.font.size = Pt(11)
        run.font.color.rgb = theme['text_secondary']
        _set_paragraph_spacing(p, 4, 0)

    _add_page_break(doc)


def _collect_headings(sections: List[Dict], depth: int = 1) -> List[Dict]:
    """递归收集所有标题，用于生成静态目录"""
    headings = []
    for sec in sections:
        if sec.get('type') == 'heading' and sec.get('title'):
            headings.append({'title': sec['title'], 'level': depth})
            headings.extend(_collect_headings(sec.get('children', []), depth + 1))
    return headings


def _build_toc(doc: Document, data: Dict, theme: Dict):
    """生成目录页：TOC 域代码 + 静态目录 fallback"""
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目 录')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = theme['primary']
    run.font.name = theme['font_heading']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), theme['font_heading'])
    _set_paragraph_spacing(p, 40, 20)

    # TOC 域代码（Word/WPS 打开会自动刷新覆盖静态内容）
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    # 静态目录内容（作为 fallback，在不支持域代码的阅读器中也能看到）
    headings = _collect_headings(data.get('sections', []))
    for h in headings:
        if h['level'] <= 3:
            indent = '    ' * (h['level'] - 1)
            run_h = p.add_run(f"{indent}{h['title']}\n")
            run_h.font.size = Pt(11 if h['level'] == 1 else 10)
            run_h.font.bold = (h['level'] == 1)
            run_h.font.color.rgb = theme['text_body']
            run_h.font.name = theme['font_body']
            run_h._element.rPr.rFonts.set(qn('w:eastAsia'), theme['font_body'])

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

    _add_page_break(doc)


def _build_heading(doc: Document, title: str, level: int, theme: Dict):
    """添加标题 + 装饰线"""
    style_name = f'Heading {level}'
    p = doc.add_heading(title, level=level)

    # 设置字体
    for run in p.runs:
        run.font.color.rgb = theme['text_dark'] if level <= 2 else theme['text_body']
        run.font.name = theme['font_heading']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), theme['font_heading'])
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)

    # H1/H2 加装饰线
    if level <= 2:
        _add_horizontal_line(doc, theme['accent'], 2.0 if level == 1 else 1.0)


def _build_paragraph(doc: Document, content: str, theme: Dict, first_indent: bool = True):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(content)
    run.font.size = Pt(10.5)
    run.font.color.rgb = theme['text_body']
    run.font.name = theme['font_body']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), theme['font_body'])

    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(p, 4, 4, 18)

    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # ~2em

    return p


def _build_quote(doc: Document, content: str, theme: Dict):
    """添加引用框（左边框 + 背景色）"""
    p = doc.add_paragraph()
    run = p.add_run(content)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = theme['text_body']
    run.font.name = theme['font_body']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), theme['font_body'])

    _set_paragraph_spacing(p, 8, 8)
    p.paragraph_format.left_indent = Cm(0.5)

    # 左边框 + 背景
    accent_hex = f'{theme["accent"][0]:02x}{theme["accent"][1]:02x}{theme["accent"][2]:02x}'
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="{accent_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # 背景色
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{theme["quote_bg"]}" w:val="clear"/>')
    pPr.append(shd)

    return p


def _build_table(doc: Document, section: Dict, theme: Dict):
    """添加数据表格"""
    headers = section.get('headers', [])
    rows = section.get('rows', [])
    cols = len(headers) if headers else (len(rows[0]) if rows else 2)

    table = doc.add_table(rows=0, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    if headers:
        hdr_row = table.add_row()
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = str(h)
            _set_cell_shading(cell, theme['table_header_bg'])
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9.5)
                    run.font.name = theme['font_body']
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), theme['font_body'])

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            # 交替行色
            if row_idx % 2 == 1:
                _set_cell_shading(cell, theme['table_alt_row'])
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = theme['text_body']
                    run.font.name = theme['font_body']
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), theme['font_body'])

    # 表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()  # 表后空行


def _build_chart(doc: Document, section: Dict, chart_paths: Dict, theme: Dict):
    """嵌入图表图片"""
    chart_id = section.get('chart_id', '')
    caption = section.get('caption', '')
    width = section.get('width_inches', 5.5)

    if chart_id not in chart_paths:
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(chart_paths[chart_id], width=Inches(width))

    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.size = Pt(9)
        run.font.color.rgb = theme['text_secondary']
        run.font.name = theme['font_body']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), theme['font_body'])
        _set_paragraph_spacing(p, 2, 8)


def _build_kpi(doc: Document, section: Dict, theme: Dict):
    """KPI 指标卡 — 用表格模拟多列布局"""
    metrics = section.get('metrics', [])
    if not metrics:
        return

    cols = len(metrics)
    table = doc.add_table(rows=3, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 去除表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    for i, m in enumerate(metrics):
        # Row 0: label
        cell = table.cell(0, i)
        cell.text = m.get('label', '')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = theme['text_secondary']

        # Row 1: value
        cell = table.cell(1, i)
        cell.text = str(m.get('value', ''))
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = theme['text_dark']

        # Row 2: change
        cell = table.cell(2, i)
        change = str(m.get('change', ''))
        cell.text = change
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
                if change.startswith('+'):
                    run.font.color.rgb = theme['positive']
                elif change.startswith('-'):
                    run.font.color.rgb = theme['negative']
                else:
                    run.font.color.rgb = theme['text_secondary']

    doc.add_paragraph()  # 后空行


def _build_section(doc: Document, section: Dict, chart_paths: Dict,
                   theme: Dict, depth: int = 1):
    """递归构建章节"""
    sec_type = section.get('type', 'heading')

    if sec_type == 'heading':
        _build_heading(doc, section.get('title', ''), depth, theme)
        content = section.get('content', '')
        if content:
            _build_paragraph(doc, content, theme, first_indent=False)

    elif sec_type == 'paragraph':
        _build_paragraph(doc, section.get('content', ''), theme)

    elif sec_type == 'quote':
        _build_quote(doc, section.get('content', ''), theme)

    elif sec_type == 'table':
        _build_table(doc, section, theme)

    elif sec_type == 'chart':
        _build_chart(doc, section, chart_paths, theme)

    elif sec_type == 'kpi':
        _build_kpi(doc, section, theme)

    elif sec_type == 'pagebreak':
        _add_page_break(doc)

    # 递归子章节
    for child in section.get('children', []):
        _build_section(doc, child, chart_paths, theme, depth + 1)


def _setup_header_footer(doc: Document, data: Dict, theme: Dict):
    """设置页眉页脚"""
    section = doc.sections[0]

    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(data.get('title', ''))
    run.font.size = Pt(8)
    run.font.color.rgb = theme['text_secondary']

    # 右侧加作者
    run2 = p.add_run('\t\t' + data.get('author', ''))
    run2.font.size = Pt(8)
    run2.font.color.rgb = theme['text_secondary']

    # 页眉下边框
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="cccccc"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # 页脚（页码）
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run('— ')
    run.font.size = Pt(8)
    run.font.color.rgb = theme['text_secondary']

    # 插入页码域
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_pg = p.add_run()
    run_pg._r.append(fldChar1)

    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_pg2 = p.add_run()
    run_pg2._r.append(instrText)

    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run_pg3 = p.add_run()
    run_pg3._r.append(fldChar2)

    run_pg4 = p.add_run('1')
    run_pg4.font.size = Pt(8)
    run_pg4.font.color.rgb = theme['text_secondary']

    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_pg5 = p.add_run()
    run_pg5._r.append(fldChar3)

    run_end = p.add_run(' —')
    run_end.font.size = Pt(8)
    run_end.font.color.rgb = theme['text_secondary']


def _build_disclaimer(doc: Document, data: Dict, theme: Dict):
    """免责声明"""
    text = data.get('disclaimer', '')
    if not text:
        return

    _add_page_break(doc)
    _build_heading(doc, '免责声明', 1, theme)

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = theme['text_secondary']
    run.font.name = theme['font_body']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), theme['font_body'])
    _set_paragraph_spacing(p, 4, 4, 14)


# ═══════════════════════════════════════════
# 主入口
# ═══════════════════════════════════════════

def render_docx(data: Dict, output: str, template: str = 'default',
                theme: str = 'cicc') -> str:
    """
    统一 DOCX 渲染入口。JSON 接口与 render_pdf 完全一致。

    Args:
        data: 报告数据字典
        output: 输出 DOCX 路径
        template: 模板名（预留）
        theme: 主题名 (cicc/ms/cms/dachen)

    Returns:
        输出文件路径
    """
    os.makedirs(os.path.dirname(output) or '.', exist_ok=True)
    th = THEMES.get(theme, THEMES['cicc'])

    doc = Document()

    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # 页眉页脚
    _setup_header_footer(doc, data, th)

    with tempfile.TemporaryDirectory(prefix='typeset_docx_') as work_dir:
        # 1. 图表
        chart_paths = _generate_charts(data, work_dir, theme)
        print(f"  Charts generated: {len(chart_paths)}")

        # 2. 封面
        _build_cover(doc, data, th)

        # 3. 目录
        if data.get('toc', True):
            _build_toc(doc, data, th)

        # 4. 正文
        for section_data in data.get('sections', []):
            _build_section(doc, section_data, chart_paths, th, depth=1)

        # 5. 免责声明
        _build_disclaimer(doc, data, th)

        # 设置 Word 打开时自动更新域（TOC 等）
        doc.settings.element.append(
            parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>')
        )

        # 保存
        doc.save(output)
        print(f"  DOCX generated: {output} ({os.path.getsize(output):,} bytes)")

    return output


if __name__ == '__main__':
    import click

    @click.command()
    @click.option('--data', required=True, help='Input JSON data file')
    @click.option('--output', required=True, help='Output DOCX path')
    @click.option('--theme', default='cicc',
                  type=click.Choice(['cicc', 'ms', 'cms', 'dachen']))
    def main(data, output, theme):
        with open(data) as f:
            d = json.load(f)
        render_docx(d, output, theme=theme)
        click.echo(f"OK: {output}")

    main()
